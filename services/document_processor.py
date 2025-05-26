from pathlib import Path
from docx import Document
from .placeholder_service import PlaceholderExtractor
import re
from datetime import datetime
import logging

# Configura o logger para esta classe
logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self, upload_folder, output_folder):
        self.templates_path = Path(upload_folder)
        self.output_path = Path(output_folder)
        self.template_placeholders = {}
    
    def get_available_templates(self):
        templates = []
        for file in self.templates_path.glob('*'):
            if file.is_file() and self._is_allowed(file.name):
                templates.append(self._get_template_info(file))
        return templates
    
    def _get_template_info(self, file):
        if file.name not in self.template_placeholders:
            if file.suffix.lower() == '.docx':
                self.template_placeholders[file.name] = PlaceholderExtractor.extract(file)
            else:
                self.template_placeholders[file.name] = self._extract_from_text(file)
        
        return {
            'filename': file.name,
            'name': file.stem,
            'path': str(file),
            'placeholders': self.template_placeholders[file.name]
        }
    
    def _is_allowed(self, filename):
        return '.' in filename and filename.rsplit('.', 1)[1].lower() in {'docx', 'doc', 'txt'}
    
    def _extract_from_text(self, file):
        with open(file, 'r', encoding='utf-8') as f:
            content = f.read()
        pattern = re.compile(r'\{\{(\w+)\}\}|\[(\w+)\]')
        return sorted({match.group(1) or match.group(2) for match in pattern.finditer(content) if match.group(1) or match.group(2)})
    
    def process_document(self, template_name, form_data, output_name):
        """Processa o documento substituindo os placeholders"""
        try:
            template_path = self.templates_path / template_name
            
            if not template_path.exists():
                raise FileNotFoundError(f"Template {template_name} não encontrado")
            
            # Garante que a pasta de saída existe
            self.output_path.mkdir(exist_ok=True)
            
            if template_path.suffix.lower() == '.docx':
                doc = self._process_docx(template_path, form_data)
                output_path = self.output_path / f"{output_name}.docx"
                doc.save(output_path)
            else:
                content = self._process_text_file(template_path, form_data)
                output_path = self.output_path / f"{output_name}.txt"
                with open(output_path, 'w', encoding='utf-8') as file:
                    file.write(content)
            
            return str(output_path)
            
        except Exception as e:
            logger.error(f"Erro ao processar documento: {str(e)}")
            raise

    def _process_docx(self, template_path, form_data):
        """Processa arquivos DOCX substituindo placeholders"""
        try:
            doc = Document(template_path)
            
            # Substitui em parágrafos
            for paragraph in doc.paragraphs:
                self._replace_in_paragraph(paragraph, form_data)
            
            # Substitui em tabelas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_in_paragraph(paragraph, form_data)
            
            return doc
            
        except ImportError:
            raise ImportError("python-docx não instalado. Use: pip install python-docx")

    def _replace_in_paragraph(self, paragraph, form_data):
        """Substitui {{CAMPO}} mantendo a semântica de chaves duplas"""
        for key, value in form_data.items():
            # Procura por {{CAMPO}} e substitui pelo valor
            placeholder = '{{' + key + '}}'
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))

    def _process_text_file(self, template_path, form_data):
        with open(template_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        for key, value in form_data.items():
            content = content.replace('{{' + key + '}}', str(value))
        
        return content