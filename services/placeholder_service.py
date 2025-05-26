from docx import Document
import re

class PlaceholderExtractor:
    @staticmethod
    def extract(docx_path):
        """Extrai placeholders de um arquivo DOCX"""
        doc = Document(docx_path)
        placeholders = set()
        pattern = re.compile(r'\{\{(\w+)\}\}|\[(\w+)\]')
        
        # Processa parágrafos
        for paragraph in doc.paragraphs:
            placeholders.update(PlaceholderExtractor._find_placeholders(paragraph.text, pattern))
        
        # Processa tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        placeholders.update(PlaceholderExtractor._find_placeholders(paragraph.text, pattern))
        
        return sorted(placeholders)
    
    @staticmethod
    def _find_placeholders(text, pattern):
        """Auxiliar para extrair placeholders de um texto"""
        matches = pattern.finditer(text)
        return {match.group(1) or match.group(2) for match in matches if match.group(1) or match.group(2)}